#!/usr/bin/env python3
"""
Generate documentation files (Excel, Word, PDF) for template
Usage: python3 generate-docs.py <template_key>
"""

import json
import sys
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF


def load_template(template_key):
    """Load template JSON"""
    json_path = f"{template_key}.template.json"
    if not Path(json_path).exists():
        print(f"❌ Template file not found: {json_path}")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def generate_excel(template, output_dir):
    """Generate Excel spreadsheet"""
    category = template['metadata'].get('category', 'General')
    template_name = template['name']
    fields = template['lists'][0]['fieldDefinitions']
    stages = template['lists'][0]['stages']

    wb = Workbook()
    ws = wb.active
    ws.title = 'Overview'

    # Header styling
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=12)

    # Template info
    ws['A1'] = 'Template'
    ws['B1'] = template_name
    ws['A1'].fill = header_fill
    ws['A1'].font = header_font

    ws['A2'] = 'Category'
    ws['B2'] = category
    ws['A2'].fill = header_fill
    ws['A2'].font = header_font

    # Fields table
    ws['A4'] = 'Field Name'
    ws['B4'] = 'Type'
    ws['C4'] = 'Required'
    for cell in ['A4', 'B4', 'C4']:
        ws[cell].font = Font(bold=True, size=11)
        ws[cell].fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    row = 5
    for field in fields:
        ws[f'A{row}'] = field['name']
        ws[f'B{row}'] = field['type']
        ws[f'C{row}'] = 'Yes' if field.get('required') else 'No'
        row += 1

    # Stages table
    ws['E4'] = 'Workflow Stage'
    ws['F4'] = 'Color'
    for cell in ['E4', 'F4']:
        ws[cell].font = Font(bold=True, size=11)
        ws[cell].fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    row = 5
    for stage in stages:
        ws[f'E{row}'] = stage['name']
        ws[f'F{row}'] = stage['color']
        # Color the cell
        fill = PatternFill(start_color=stage['color'][1:], end_color=stage['color'][1:], fill_type='solid')
        ws[f'F{row}'].fill = fill
        row += 1

    # Auto-width columns
    for col in ['A', 'B', 'C', 'E', 'F']:
        ws.column_dimensions[col].width = 20

    # Category-specific sheet
    if category in ['Human Resources', 'HR', 'Recruitment']:
        ws2 = wb.create_sheet('Evaluation Form')
        ws2['A1'] = 'Candidate Evaluation Scorecard'
        ws2['A1'].font = Font(bold=True, size=14)
        ws2['A3'] = 'Criteria'
        ws2['B3'] = 'Score (1-5)'
        ws2['C3'] = 'Notes'
        criteria = ['Technical Skills', 'Communication', 'Culture Fit', 'Experience', 'Problem Solving']
        for i, crit in enumerate(criteria, 4):
            ws2[f'A{i}'] = crit

    elif category in ['Development', 'Software', 'Tech']:
        ws2 = wb.create_sheet('Bug Tracking')
        ws2['A1'] = 'Bug Severity Matrix'
        ws2['A1'].font = Font(bold=True, size=14)
        ws2['A3'] = 'Priority'
        ws2['B3'] = 'Impact'
        ws2['C3'] = 'Urgency'
        ws2['D3'] = 'Action'

    elif category in ['Sales', 'Marketing']:
        ws2 = wb.create_sheet('Pipeline Tracking')
        ws2['A1'] = 'Sales Pipeline Dashboard'
        ws2['A1'].font = Font(bold=True, size=14)
        ws2['A3'] = 'Stage'
        ws2['B3'] = 'Count'
        ws2['C3'] = 'Value'
        for i, stage in enumerate(stages, 4):
            ws2[f'A{i}'] = stage['name']

    wb.save(output_dir / 'template_data.xlsx')
    print(f"✅ Excel: {output_dir}/template_data.xlsx")


def generate_word(template, output_dir):
    """Generate Word document"""
    category = template['metadata'].get('category', 'General')
    template_name = template['name']
    description = template['description']
    fields = template['lists'][0]['fieldDefinitions']
    stages = template['lists'][0]['stages']

    doc = Document()

    # Title
    doc.add_heading(template_name, 0)
    doc.add_paragraph(description)

    # Workflow Stages
    doc.add_heading('Workflow Stages', 1)
    for i, stage in enumerate(stages, 1):
        doc.add_paragraph(f"{i}. {stage['name']}", style='List Number')

    # Fields Reference
    doc.add_heading('Fields Reference', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Required'

    for field in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['name']
        row_cells[1].text = field['type']
        row_cells[2].text = '✓' if field.get('required') else ''

    # Category-specific content
    doc.add_page_break()

    if category in ['Human Resources', 'HR', 'Recruitment']:
        doc.add_heading('Interview Guide', 1)
        doc.add_paragraph('Use this template for structured interviews.')
        doc.add_heading('Preparation', 2)
        doc.add_paragraph('• Review CV 15 minutes before', style='List Bullet')
        doc.add_paragraph('• Prepare role-specific questions', style='List Bullet')
        doc.add_paragraph('• Check meeting room/link', style='List Bullet')

    elif category in ['Development', 'Software', 'Tech']:
        doc.add_heading('Bug Report Template', 1)
        doc.add_paragraph('Use this template when reporting bugs.')
        doc.add_heading('Required Information', 2)
        doc.add_paragraph('• Clear, descriptive title', style='List Bullet')
        doc.add_paragraph('• Steps to reproduce', style='List Bullet')
        doc.add_paragraph('• Expected vs actual behavior', style='List Bullet')
        doc.add_paragraph('• Environment details', style='List Bullet')

    elif category in ['Sales', 'Marketing']:
        doc.add_heading('Lead Qualification Guide', 1)
        doc.add_paragraph('Use this guide to qualify leads.')
        doc.add_heading('Qualification Criteria', 2)
        doc.add_paragraph('• Budget confirmed', style='List Bullet')
        doc.add_paragraph('• Authority identified', style='List Bullet')
        doc.add_paragraph('• Need validated', style='List Bullet')
        doc.add_paragraph('• Timeline established', style='List Bullet')

    doc.save(output_dir / 'template_overview.docx')
    print(f"✅ Word: {output_dir}/template_overview.docx")


def generate_pdf(template, output_dir):
    """Generate PDF guide"""
    template_name = template['name']
    description = template['description']
    category = template['metadata'].get('category', 'General')
    fields = template['lists'][0]['fieldDefinitions']
    stages = template['lists'][0]['stages']

    pdf = FPDF()
    pdf.add_page()

    # Header
    pdf.set_font('Arial', 'B', 18)
    pdf.cell(0, 15, template_name, ln=True, align='C')
    pdf.set_font('Arial', 'I', 10)
    pdf.cell(0, 8, 'Quick Reference Guide', ln=True, align='C')
    pdf.ln(5)

    # Description
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 8, description)
    pdf.ln(5)

    # Category
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(40, 8, 'Category: ', 0)
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 8, category, ln=True)
    pdf.ln(5)

    # Workflow Stages
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'Workflow Stages', ln=True)
    pdf.set_font('Arial', '', 11)

    stage_flow = ' → '.join([s['name'] for s in stages])
    pdf.multi_cell(0, 8, stage_flow)
    pdf.ln(5)

    # Fields
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'Fields', ln=True)
    pdf.set_font('Arial', '', 11)

    for field in fields:
        required = ' (Required)' if field.get('required') else ''
        pdf.cell(0, 8, f"• {field['name']} - {field['type']}{required}", ln=True)

    pdf.ln(5)

    # Getting Started
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'Getting Started', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.multi_cell(0, 8, '1. Import this template into your Priovs system\n2. Customize fields as needed\n3. Start tracking your workflow\n4. Move items through stages as work progresses')

    pdf.output(str(output_dir / 'template_guide.pdf'))
    print(f"✅ PDF: {output_dir}/template_guide.pdf")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 generate-docs.py <template_key>")
        sys.exit(1)

    template_key = sys.argv[1]

    # Load template
    print(f"📄 Loading template: {template_key}")
    template = load_template(template_key)

    # Create docs directory
    output_dir = Path('docs')
    output_dir.mkdir(exist_ok=True)
    print(f"📁 Output directory: {output_dir}")

    # Generate files
    print("\n🔨 Generating documentation files...")
    generate_excel(template, output_dir)
    generate_word(template, output_dir)
    generate_pdf(template, output_dir)

    print("\n✅ All documentation files created successfully!")
    print(f"   Total: 3 files in {output_dir}/")


if __name__ == '__main__':
    main()
